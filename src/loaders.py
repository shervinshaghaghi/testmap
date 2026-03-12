# src/loaders.py
from pathlib import Path
from docx import Document
from docx.shared import Pt
import pdfplumber

def clamp(text: str, max_chars: int) -> str:
    t = (text or "")
    if len(t) <= max_chars:
        return t
    return t[:max_chars] + "\n...[TRUNCATED]...\n"

def load_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    cells.append(ct)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)

def load_pdf_text(path: Path) -> str:
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = (page.extract_text() or "").strip()
            if t:
                parts.append(t)
    return "\n\n".join(parts)

def load_student_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == ".docx":
        return load_docx_text(path)
    if ext == ".pdf":
        return load_pdf_text(path)
    raise ValueError(f"Unsupported student file type: {ext}")