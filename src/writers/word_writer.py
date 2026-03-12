# src/writers/word_writer.py
from pathlib import Path
from docx import Document
from docx.shared import Pt


def write_result_docx(result: dict, out_path: Path) -> None:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("LLM Review.", level=2)

    # ---- mapping table ----
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Student"
    hdr[1].text = "CaseID"
    hdr[2].text = "RequirementID"
    hdr[3].text = "Rationale (short)"

    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    student = (result.get("student") or "").strip() or "STUDENT"
    student_cell_text = "\n".join([p for p in student.split() if p]) or student

    rows = result.get("rows") or []
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = student_cell_text
        cells[1].text = str(r.get("case_id", "") or "")
        cells[2].text = str(r.get("requirement_id", "") or "")
        cells[3].text = str(r.get("rationale", "") or "")

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # ---- OOB/MIX details + problematic students (as in your prompt/sample) ----
    notes = result.get("notes") or {}
    oob_or_mix = notes.get("oob_or_mix") or []
    problematic = notes.get("problematic_students")
    if problematic is None or str(problematic).strip() == "":
        problematic = "None"

    doc.add_paragraph("")  # spacing
    doc.add_heading("OOB (and MIX) details", level=3)

    if oob_or_mix:
        for item in oob_or_mix:
            cid = str(item.get("case_id", "") or "").strip()
            detail = str(item.get("detail", "") or "").strip()
            if cid and detail:
                doc.add_paragraph(f"{cid} — {detail}", style="List Bullet")
    else:
        doc.add_paragraph("None.")

    doc.add_paragraph(f"Problematic students (hard to interpret): {problematic}.")

    doc.save(str(out_path))